from docx import Document


TEMPLATE = "picur_template.docx"
OUTPUT = "Anexo1-Formato inscripción de proyectos (PICUR)-GeoGuardian_AI-diligenciado.docx"


def set_cell(cell, text):
    cell.text = text


def main():
    doc = Document(TEMPLATE)

    student_rows = [
        ("PENDIENTE - Nombre estudiante 1", "PENDIENTE", "PENDIENTE"),
        ("PENDIENTE - Nombre estudiante 2", "PENDIENTE", "PENDIENTE"),
        ("PENDIENTE - Nombre estudiante 3 (opcional)", "PENDIENTE", "PENDIENTE"),
    ]

    project_title = "GeoGuardian AI: sistema prototipo de monitoreo preventivo para zonas con riesgo geológico mediante aplicación móvil, sensores IoT e inteligencia artificial"

    introduction = (
        "GeoGuardian AI es una propuesta tecnológica orientada a apoyar la detección temprana de condiciones asociadas "
        "a derrumbes, deslizamientos e inestabilidad del terreno. El proyecto integra una aplicación móvil desarrollada "
        "en Flutter, visualización geográfica, reportes ciudadanos, sensores IoT proyectados y análisis asistido por "
        "inteligencia artificial. Su propósito es centralizar información ambiental y geotécnica básica para que usuarios "
        "de zonas vulnerables puedan consultar alertas, revisar historial de eventos y recibir recomendaciones preventivas "
        "en una interfaz clara y accesible."
    )

    problem = (
        "En comunidades ubicadas en laderas, zonas de alta pendiente o sectores afectados por lluvias intensas, la falta "
        "de información oportuna puede dificultar la identificación de señales tempranas de riesgo. Variables como humedad "
        "del suelo, lluvia acumulada, inclinación del terreno, cambios visuales y reportes comunitarios suelen observarse "
        "de manera aislada. Esta situación limita la toma de decisiones preventivas y puede aumentar la exposición de las "
        "personas ante eventos de remoción en masa. Se plantea como hipótesis que una herramienta móvil que centralice "
        "sensores, mapas, historial y apoyo de IA puede mejorar la comprensión del riesgo y fortalecer la prevención."
    )

    justification = (
        "La investigación es pertinente porque aborda una necesidad social relacionada con la gestión preventiva del riesgo "
        "geológico. Su relevancia radica en acercar tecnologías como IoT, geolocalización e inteligencia artificial a usuarios "
        "no especializados, mediante una aplicación que organiza datos y entrega mensajes accionables. El impacto esperado "
        "incluye mayor conciencia comunitaria, registro ordenado de eventos, apoyo a la observación del terreno y una base "
        "técnica para futuras integraciones con sensores reales, backend, modelos predictivos y sistemas de alerta temprana."
    )

    objectives = (
        "Objetivo general: Desarrollar un prototipo de sistema de monitoreo preventivo que permita visualizar el estado del "
        "terreno, consultar alertas, revisar historial de eventos, reportar zonas de riesgo y recibir apoyo de IA para "
        "interpretar condiciones del entorno.\n\n"
        "Objetivos específicos:\n"
        "1. Diseñar una aplicación móvil con autenticación, panel principal, mapa de riesgo, historial, alertas y ajustes.\n"
        "2. Simular mediciones de humedad, lluvia, inclinación y temperatura para validar la experiencia de monitoreo.\n"
        "3. Incorporar visualización geográfica de zonas seguras, de precaución y de riesgo alto.\n"
        "4. Permitir el reporte de condiciones observadas por el usuario en una zona determinada.\n"
        "5. Simular análisis de imagen y asistencia textual con IA para orientar recomendaciones preventivas.\n"
        "6. Definir una arquitectura preparada para integrar backend, sensores ESP32 y modelos reales de inteligencia artificial."
    )

    theory = (
        "El proyecto se fundamenta en la gestión del riesgo de desastres, los sistemas de alerta temprana, el monitoreo "
        "geológico y el uso de tecnologías digitales para apoyar la prevención. La UNDRR define los sistemas de alerta "
        "temprana como procesos integrados de monitoreo, evaluación del riesgo, comunicación y preparación para reducir "
        "daños antes de un evento peligroso. En el caso de deslizamientos, el monitoreo continuo permite reconocer cambios "
        "en variables físicas y ambientales que pueden anteceder movimientos del terreno, como lluvias intensas, saturación "
        "del suelo e inclinación. La literatura técnica también resalta el valor de combinar sensores, análisis espacial, "
        "modelos de predicción y comunicación clara con la comunidad.\n\n"
        "Desde la perspectiva tecnológica, IoT permite capturar datos mediante nodos de sensado conectados, mientras que una "
        "aplicación móvil facilita que el usuario consulte el estado del sistema en tiempo real o semirreal. La inteligencia "
        "artificial puede apoyar la clasificación de imágenes, la estimación de riesgo y la generación de recomendaciones, "
        "siempre que sus resultados sean validados por criterios técnicos. GeoGuardian AI adopta estos principios como un "
        "prototipo educativo y funcional que organiza información de sensores simulados, mapa, historial, alertas y asistencia "
        "IA en una arquitectura escalable."
    )

    methodology = (
        "La investigación se plantea como aplicada y de desarrollo tecnológico, con enfoque descriptivo y experimental a nivel "
        "de prototipo. El diseño contempla cuatro capas: aplicación móvil, backend/API futuro, dispositivos IoT ESP32 e "
        "inteligencia artificial. La población objetivo corresponde a habitantes de zonas vulnerables, equipos comunitarios "
        "de prevención, estudiantes e investigadores interesados en monitoreo ambiental.\n\n"
        "Las técnicas de recolección previstas incluyen revisión documental, observación de necesidades de usuario, pruebas "
        "funcionales del prototipo y, en fases posteriores, captura de datos desde sensores reales de humedad, lluvia, "
        "inclinación y temperatura. Actualmente el prototipo usa datos simulados para validar navegación, visualización y "
        "flujo de interacción. Como aspectos éticos, se considera la protección de datos personales y de ubicación, la "
        "necesidad de no presentar recomendaciones automatizadas como diagnósticos definitivos y la validación técnica antes "
        "de usar el sistema para decisiones de seguridad pública."
    )

    results = (
        "Al encontrarse en etapa de propuesta/prototipo, los resultados esperados son: una aplicación móvil funcional con "
        "inicio de sesión, registro, verificación, recuperación de contraseña, dashboard, tarjetas de sensores, mapa de "
        "riesgo, reporte de zona, cámara IA simulada, historial de eventos, alertas, perfil, seguridad y asistente IA. "
        "También se espera contar con una arquitectura documentada para integrar backend, ESP32 e IA real. En términos de "
        "impacto, se espera que el prototipo permita demostrar cómo la centralización de datos y alertas puede apoyar la "
        "prevención comunitaria y servir como base para futuras pruebas con datos reales."
    )

    conclusions = (
        "GeoGuardian AI propone una solución preventiva, escalable y orientada al usuario para organizar información sobre "
        "riesgo geológico. El prototipo demuestra que una aplicación móvil puede integrar visualización de sensores, mapas, "
        "reportes, historial y asistencia IA en una experiencia sencilla. Aunque aún requiere conexión con sensores reales, "
        "backend seguro y validación técnica, el proyecto constituye una base viable para fortalecer procesos de monitoreo "
        "temprano y educación comunitaria frente a deslizamientos."
    )

    sources = (
        "1. United Nations Office for Disaster Risk Reduction. (2017). Sendai Framework Terminology on Disaster Risk Reduction: Early warning system. https://www.undrr.org/terminology/early-warning-system\n"
        "2. United Nations Office for Disaster Risk Reduction. (2015). Sendai Framework for Disaster Risk Reduction 2015-2030. https://www.undrr.org/publication/sendai-framework-disaster-risk-reduction-2015-2030\n"
        "3. U.S. Geological Survey. (s. f.). Landslide Hazards Program: Monitoring. https://www.usgs.gov/programs/landslide-hazards/monitoring\n"
        "4. U.S. Geological Survey. (s. f.). Landslide basics. https://www.usgs.gov/programs/landslide-hazards/science/landslide-basics\n"
        "5. U.S. Geological Survey. (1999). Real-time monitoring of active landslides. https://pubs.usgs.gov/fs/fs-091-99/fs091_99.html\n"
        "6. World Meteorological Organization. (2022). Early warning and early action. https://wmo.int/about-us/world-meteorological-day/wmd-2022\n"
        "7. NASA Goddard Space Flight Center. (s. f.). Landslide Hazard Assessment for Situational Awareness (LHASA). https://github.com/nasa/LHASA\n"
        "8. Highland, L. M., & Bobrowsky, P. (2008). The landslide handbook: A guide to understanding landslides. U.S. Geological Survey Circular 1325.\n"
        "9. Flutter. (s. f.). Flutter documentation. https://docs.flutter.dev/\n"
        "10. OpenStreetMap Foundation. (s. f.). OpenStreetMap. https://www.openstreetmap.org/\n"
        "11. Documentación interna del proyecto GeoGuardian AI: visión general, arquitectura, aplicación Flutter y estado actual."
    )

    cronograma = [
        ("Revisión documental y definición del problema", "Agosto 2026", "Agosto 2026", "Equipo del proyecto", "Documentos, repositorio, bibliografía", "Base conceptual y alcance inicial"),
        ("Diseño de arquitectura y experiencia móvil", "Agosto 2026", "Septiembre 2026", "Equipo del proyecto", "Flutter, Figma, documentación técnica", "Definición de pantallas y flujo"),
        ("Construcción del prototipo Flutter", "Septiembre 2026", "Octubre 2026", "Equipo del proyecto", "Flutter SDK, VS Code, emulador/dispositivo", "App con datos simulados"),
        ("Simulación de sensores, mapa, alertas e IA", "Octubre 2026", "Noviembre 2026", "Equipo del proyecto", "OpenStreetMap, servicios locales, datos demo", "Validación funcional del prototipo"),
        ("Pruebas, documentación y socialización", "Noviembre 2026", "Noviembre 2026", "Equipo del proyecto y asesores", "Guías, pruebas, presentación", "Entrega del prototipo y mejoras pendientes"),
    ]

    t0 = doc.tables[0]
    for idx, (name, cedula, email) in enumerate(student_rows, start=2):
        set_cell(t0.rows[idx].cells[0], name)
        set_cell(t0.rows[idx].cells[3], cedula)
        set_cell(t0.rows[idx].cells[4], email)

    set_cell(t0.rows[6].cells[1], "PENDIENTE - Asesor temático")
    set_cell(t0.rows[6].cells[3], "PENDIENTE")
    set_cell(t0.rows[6].cells[4], "PENDIENTE")
    set_cell(t0.rows[7].cells[1], "PENDIENTE - Asesor metodológico")
    set_cell(t0.rows[7].cells[3], "PENDIENTE")
    set_cell(t0.rows[7].cells[4], "PENDIENTE")

    t1 = doc.tables[1]
    section_values = {
        1: project_title,
        3: introduction,
        5: problem,
        7: justification,
        9: objectives,
        11: theory,
        13: methodology,
        15: results,
        17: conclusions,
        19: sources,
    }
    for row_idx, text in section_values.items():
        set_cell(t1.rows[row_idx].cells[0], text)

    t2 = doc.tables[2]
    for row_idx, values in enumerate(cronograma, start=2):
        for col_idx, value in enumerate(values):
            set_cell(t2.rows[row_idx].cells[col_idx], value)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
